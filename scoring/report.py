import io
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Word helpers ──────────────────────────────────────────────────────────────

def _bidi(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    b = OxmlElement('w:bidi')
    b.set(qn('w:val'), '1')
    pPr.append(b)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _font(run, bold=False, size=11, color=None):
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def _heading(doc, text, level=2):
    p = doc.add_paragraph()
    _bidi(p)
    run = p.add_run(text)
    _font(run, bold=True, size=13 if level == 1 else 11, color=(44, 95, 138))
    return p


def _para(doc, text):
    p = doc.add_paragraph()
    _bidi(p)
    run = p.add_run(text)
    _font(run)
    return p


def _tbl_bidi(table):
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    bv = OxmlElement('w:bidiVisual')
    tblPr.append(bv)


def _cell_write(cell, text, bold=False, bg=None):
    cell.text = ''
    p = cell.paragraphs[0]
    _bidi(p)
    run = p.add_run(str(text) if text is not None else '—')
    _font(run, bold=bold)
    if bg:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  bg)
        tcPr.append(shd)


def _add_table(doc, headers, rows, header_bg='2C5F8A'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    _tbl_bidi(table)

    hrow = table.rows[0]
    for i, h in enumerate(headers):
        _cell_write(hrow.cells[i], h, bold=True, bg=header_bg)

    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = 'F0F4F8' if ri % 2 == 0 else None
        for ci, val in enumerate(row_data):
            _cell_write(row.cells[ci], val, bg=bg)

    return table


# ── Word document ─────────────────────────────────────────────────────────────

def generate_docx(results: dict) -> bytes:
    doc   = Document()
    child = results['child']

    # page margins
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.right_margin  = Cm(2)
        section.left_margin   = Cm(2)

    # title
    t = doc.add_paragraph()
    _bidi(t)
    _font(t.add_run('דוח אבחון WISC-IV'), bold=True, size=16, color=(44, 95, 138))

    doc.add_paragraph()

    # child info block
    _heading(doc, 'פרטי הנבדק')
    info = [
        ('שם',               child.get('name', '')),
        ('תאריך לידה',       child.get('dob', '')),
        ('תאריך בדיקה',      child.get('exam_date', '')),
        ('גיל בבדיקה',       child.get('age_str', '')),
        ('בוחן',             child.get('examiner', '')),
        ('מסגרת חינוכית',    child.get('school', '')),
        ('כיתה',             child.get('grade', '')),
        ('סיבת הפניה',       child.get('reason', '')),
    ]
    for label, val in info:
        if val:
            _para(doc, f'{label}: {val}')

    doc.add_paragraph()

    # ── subtest table ──
    _heading(doc, 'ציוני מבחנים')

    st_headers = ['מבחן', 'ציון גולמי', 'תת-ציונים', 'ציון תקני', 'תיאור']
    st_rows    = []
    for code in ['BD', 'SI', 'DS', 'PCn', 'CD', 'VC', 'LN', 'MR', 'CO', 'SS',
                 'PCm', 'CA', 'IN', 'AR', 'WR']:
        e = results['subtests'].get(code)
        if not e:
            continue
        detail = ''
        if 'sub_a' in e and e['sub_a'] is not None:
            diff_sign = '+' if e.get('sub_diff', 0) >= 0 else ''
            detail = (
                f"{e['label_a']}: {e['sub_a']} | "
                f"{e['label_b']}: {e['sub_b']} | "
                f"הפרש: {diff_sign}{e.get('sub_diff', '—')}"
            )
        st_rows.append([
            e['name'],
            e.get('raw', '—'),
            detail or '—',
            e.get('scaled', '—'),
            e.get('qualitative', '—'),
        ])
    _add_table(doc, st_headers, st_rows)

    doc.add_paragraph()

    # ── composite table ──
    _heading(doc, 'מדדים מורכבים')

    comp_headers = ['מדד', 'ציון', 'אחוזון', 'ר"ב 90%', 'ר"ב 95%', 'תיאור']
    comp_rows    = []
    for idx in ['VCI', 'PRI', 'WMI', 'PSI', 'FSIQ']:
        c   = results['composites'].get(idx, {})
        ci90 = (f"{c.get('ci90_lo', '—')}–{c.get('ci90_hi', '—')}"
                if c.get('score') else '—')
        ci95 = (f"{c.get('ci95_lo', '—')}–{c.get('ci95_hi', '—')}"
                if c.get('score') else '—')
        comp_rows.append([
            c.get('name', idx),
            c.get('score', '—'),
            c.get('percentile', c.get('error', '—')),
            ci90, ci95,
            c.get('qualitative', '—'),
        ])
    _add_table(doc, comp_headers, comp_rows)

    doc.add_paragraph()

    # ── discrepancy table ──
    disc = results.get('discrepancies', [])
    if disc:
        _heading(doc, 'הפרשים בין מדדים')
        d_headers = ['מדד א׳', 'ציון א׳', 'מדד ב׳', 'ציון ב׳', 'הפרש', 'משמעותי?']
        d_rows    = [
            [d['name_a'], d['score_a'], d['name_b'], d['score_b'],
             d['abs_diff'], 'כן ✓' if d['significant'] else 'לא']
            for d in disc
        ]
        _add_table(doc, d_headers, d_rows)
        doc.add_paragraph()

    # ── interpretation ──
    interp = results.get('interpretation', {})
    if interp:
        _heading(doc, 'פרשנות')
        for key in ['fsiq', 'vci', 'pri', 'wmi', 'psi',
                    'discrepancies', 'strengths', 'weaknesses']:
            txt = interp.get(key)
            if txt:
                _para(doc, txt)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ── PDF via weasyprint ────────────────────────────────────────────────────────

def generate_pdf(results: dict, templates_dir: str) -> bytes:
    from jinja2 import Environment, FileSystemLoader
    import weasyprint

    env      = Environment(loader=FileSystemLoader(templates_dir))
    template = env.get_template('report_print.html')
    html_str = template.render(results=results)

    return weasyprint.HTML(string=html_str).write_pdf()
